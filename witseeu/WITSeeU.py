# WITSeeU.py
from __future__ import unicode_literals

# Flask & Web Utilities
from flask import (
    Flask, render_template, jsonify, request,
    json, session
)

# Standard Libraries
import requests
from datetime import datetime
import pytz
import os
import re
import json_repair
import tiktoken

# Document Processing
import PyPDF2
import docx

# AI / LLM
from openai import OpenAI

# CORS
from flask_cors import CORS

# Prompt Generator
from prompt_generator import PromptGenerator

# --- 1. Flask App 初始化與設定 ---
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', os.urandom(24))

# 建立 V1.0 實例 (一般規則)
pg_v1 = PromptGenerator(app_version='1.1')

# 建立 V2.0 實例 (RAG 規則)
pg_v2 = PromptGenerator(app_version='2.0')

# App Configurations
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['LOG_FOLDER'] = 'log'
app.config['TIME_STAMP'] = datetime.now().strftime("%m%d%H%M%S")

# 時區
tz_gmt_plus_8 = pytz.timezone('Asia/Taipei')

# OpenAI Client
clientAI = OpenAI(api_key=os.environ.get('OPENAI_KEY'))

# RAGFlow 設定
RAGFLOW_API_KEY = os.environ.get('RAGFLOW_API_KEY')
RAGFLOW_API_URL = os.environ.get('RAGFLOW_API_URL')

# CORS
CORS(app, resources={
    r"/WITSeeU": {
        "origins": "*"
    }
})

# AI Model
DEFAULT_AI_MODEL = 'gpt-4o-mini'

# Processing Mode
PROC_MODE = 'debug'

# tiktoken
encoding = tiktoken.encoding_for_model(DEFAULT_AI_MODEL)

def wprint(t):
    if PROC_MODE == 'debug':
        print(t)

def reset_session():
    session.clear()

def report_usage(action, response):
    inToken = 0
    outToken = 0

    if action == 'inToken':
        try:
            text = json.dumps(response, ensure_ascii=False)
            inToken = len(encoding.encode(text))
        except Exception as e:
            print(f"[report_usage] in token 計算失敗: {e}")

    elif action == 'outToken':
        try:
            text = response.text if hasattr(response, "text") else str(response)
            outToken = len(encoding.encode(text))
        except Exception as e:
            print(f"[report_usage] out token 計算失敗: {e}")

    else:
        try:
            inToken = int(response.usage.prompt_tokens)
            outToken = int(response.usage.completion_tokens)
        except Exception as e:
            print(f"[report_usage] LLM token 計算失敗: {e}")

    # 總 token
    toToken = inToken + outToken
    session['input_tokens'] = session.get('input_tokens', 0) + inToken
    session['output_tokens'] = session.get('output_tokens', 0) + outToken
    session['total_tokens'] = session.get('total_tokens', 0) + toToken

    wprint(f"{action} : ")
    wprint(f"input tokens = {inToken}")
    wprint(f"output tokens = {outToken}")
    wprint(f"total tokens = {toToken}\n")

def save_html_to_file(html_content):
    filename = f"WITSeeU_output_{session.get('curDateTime')}.html"
    file_path = os.path.join(app.config['LOG_FOLDER'], filename)
    if not os.path.exists(app.config['LOG_FOLDER']):
        os.makedirs(app.config['LOG_FOLDER'])
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(html_content)
    wprint(f"HTML 檔案已儲存：{file_path}")

def save_to_file(filename, content):
    file_path = os.path.join(app.config['LOG_FOLDER'], session.get('curDateTime') + '_' + filename)
    if not os.path.exists(app.config['LOG_FOLDER']):
        os.makedirs(app.config['LOG_FOLDER'])
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(str(content))
    wprint(f"檔案已儲存：{file_path}")

def extract_text_from_pdf(pdf_path):
    wprint("解析 PDF 文件並提取文本")
    text = ""
    with open(pdf_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx_table(docx_path):
    doc = docx.Document(docx_path)
    text = []
    seen_texts = set()
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                cell_content = "\n".join(cell_paragraphs)
                if cell_content and cell_content not in seen_texts:
                    seen_texts.add(cell_content)
                    row_text.append(cell_content)
                else:
                    row_text.append("")
            text.append('\t'.join(row_text))
    if not text:
        text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return text

def extract_text(file_path):
    wprint("根據文件類型選擇適當的解析方法")
    if not file_path:
        return None
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx_table(file_path)
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    else:
        return None

def fix_json(text):
    # 處理：缺少引號、缺少括號、巢狀結構、雜亂文字
    decoded_object = json_repair.repair_json(text, return_objects=True)
    return json.dumps(decoded_object, indent=4, ensure_ascii=False)

def gen_structured_jd(jd_text):
    system_prompt = None
    user_prompt = None
    messages = None
    content = None

    if session['appVersion'] == '2.0':
        system_prompt = pg_v2.get_jd_system_prompt()
        user_prompt = pg_v2.get_jd_user_prompt(jd_text)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        content = callRAG(messages)

    else :
        system_prompt = pg_v1.get_jd_system_prompt()
        user_prompt = pg_v1.get_jd_user_prompt(jd_text)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        content =  callLLM(messages)

    return fix_json(content)

def gen_match_score(jd, resume):
    system_prompt = None
    user_prompt = None
    messages = None
    content = None

    if session['appVersion'] == '2.0':
        system_prompt = pg_v2.get_match_score_system_prompt()
        user_prompt = pg_v2.get_match_score_user_prompt(jd, resume)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        content = callRAG(messages)

    else :
        system_prompt = pg_v1.get_match_score_system_prompt()
        user_prompt = pg_v1.get_match_score_user_prompt(jd, resume)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        content =  callLLM(messages)

    return fix_json(content)

def gen_match_html(jd, resume):
    system_prompt = None
    user_prompt = None
    messages = None
    content = None

    if session['appVersion'] == '2.0':
        system_prompt = pg_v2.get_match_html_system_prompt()
        user_prompt = pg_v2.get_match_html_user_prompt(session.get('structured_jd'), resume)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        content = callRAG(messages)

        html_match = re.search(r"<\s*HTML_REPORT\s*>(.*?)</\s*HTML_REPORT\s*>", content, re.DOTALL | re.IGNORECASE)
        if html_match:
            content = html_match.group(1).strip()
            wprint("INFO: 成功解析 HTML 分析報告。")
        else:
            error_msg = "回傳內容中未找到 <HTML_REPORT> 標籤 (或其格式不符)。"
            wprint(f"ERROR: {error_msg}")

    else :
        system_prompt = pg_v1.get_match_html_system_prompt()
        user_prompt = pg_v1.get_match_html_user_prompt(session.get('structured_jd'), resume)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        content =  callLLM(messages)

    html_match = re.search(r"<\s*HTML_REPORT\s*>(.*?)</\s*HTML_REPORT\s*>", content, re.DOTALL | re.IGNORECASE)
    if html_match:
        content = html_match.group(1).strip()
        wprint("INFO: 成功解析 HTML 分析報告。")
    else:
        error_msg = "回傳內容中未找到 <HTML_REPORT> 標籤 (或其格式不符)。"
        wprint(f"ERROR: {error_msg}")

    return content

def callLLM(messages):
    response = clientAI.chat.completions.create(
        model=session.get('aiModel', DEFAULT_AI_MODEL),
        temperature=0.7,
        max_tokens=2048,
        messages=messages
    )

    # 計算費用
    report_usage("inToken", messages)
    report_usage("outToken", response)

    content = response.choices[0].message.content
    return content

def callRAG(messages):
    payload = {
        "model": session.get('aiModel', "gpt-4.1-mini"),
        "stream": False,
        "max_tokens": 2048,
        "messages": messages
    }
    headers = {
        "Authorization": f"Bearer {RAGFLOW_API_KEY}",
        "Content-Type": "application/json"
    }

    response = requests.post(
        RAGFLOW_API_URL,
        headers=headers,
        json=payload
    )

    # 計算費用
    report_usage("inToken", messages)
    report_usage("outToken", response)

    response.raise_for_status()
    data = response.json()
    content = "".join([choice["message"]["content"] for choice in data.get("choices", [])])
    return content

@app.route('/WITSeeU', methods=['GET', 'POST'])
def WITSeeU():
    if request.method == 'GET':
        return jsonify({"status": "success", "message": "Service is running. Please use POST to upload files."})

    try:
        # 1. 初始化 Session 與環境變數
        reset_session()
        session['curDateTime'] = datetime.now(tz_gmt_plus_8).strftime("%y%m%d%H%M%S")

        # 2. 接收參數
        appVersion = request.form.get('appVersion', '1.1')
        session['appVersion'] = request.form.get('appVersion', '1.1')
        session['aiModel'] = request.form.get('aiModel', 'gpt-4o-mini')

        # 3. 檔案處理
        jd_file = request.files.get('jd_file')
        resume_file = request.files.get('resume_file')
        jd_path = None
        resume_path = None

        if jd_file and jd_file.filename:
            jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_file.filename)
            jd_file.save(jd_path)

        if resume_file and resume_file.filename:
            resume_path = os.path.join(app.config['UPLOAD_FOLDER'], resume_file.filename)
            resume_file.save(resume_path)

        # 4. 文字萃取
        jd_text = extract_text(jd_path)
        resume_text = extract_text(resume_path)

        # 5. 格式檢查 (若萃取失敗或無內容)
        if not jd_text or not resume_text:
            wprint("文件格式不支援或內容為空")
            return jsonify({
                "status": "error",
                "message": "文件格式不支援或無法讀取，請上傳有效的 docx / txt 檔！"
            }), 400

        # 6. 根據版本執行 AI 匹配
        # 6.1) 產生結構化 jd
        structured_jd = gen_structured_jd(jd_text)

        # 6.2) 與resume匹配, 產生匹配結果
        match_score = gen_match_score(structured_jd, resume_text)

        # 6.3) 依比對結果, 產生匹配詳細報告(HTML)
        match_html = gen_match_html(structured_jd, match_score)

        # 7. 計算 Token 費用
        input_tokens = session.get('input_tokens', 0)
        output_tokens = session.get('output_tokens', 0)
        total_tokens = session.get('total_tokens', 0)
        total_cost = (input_tokens * 0.001 * 0.00015) + (output_tokens * 0.001 * 0.0006) # 費用公式 (gpt-4o-mini)

        # 8. 組裝回傳 JSON
        response_data = {
            "status": "success",
            "answer": match_html,
            "result": match_score,
            "info": {
                "aiModel": session.get('aiModel'),
                "appVersion": appVersion,
                "jd_file": jd_file.filename if jd_file else "",
                "resume_file": resume_file.filename if resume_file else ""
            },
            "usage": {
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_tokens": total_tokens,
                "total_cost": f"{total_cost:.6f}"
            }
        }

        return jsonify(response_data)

    except Exception as e:
        wprint(f"System Error: {e}")
        save_to_file("error", f"{e}")
        return jsonify({
            "status": "error",
            "message": f"系統發生內部錯誤: {str(e)}"
        }), 500

# --- 4. 入口 ---
if __name__ == "__main__":
    wprint("main")
    # 確保上傳目錄存在
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run()
